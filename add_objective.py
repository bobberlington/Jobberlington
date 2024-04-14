from docx import Document
from docx.shared import Pt
from rewrite_resume import rewrite_resume
from credentials import resume
from cover_letter import write_cover_letter
description = """
Job ID 17612
The Space Dynamics Laboratory (SDL) is searching for a software engineer to join our team. We are global leaders in creating frameworks that facilitate high-speed image processing and the distribution of image products. Our team utilizes an iterative software development approach, delivering solutions to our clients at regular intervals. Candidates should have familiarity with one or more of the following areas: software design, APIs, data visualization, containerization, system integration, and database design.
This position is located onsite in Logan, UT.

Required Qualifications:

    Completed bachelor's degree in computer science, computer engineering, electrical engineering, or a related field by the end of 2024
    Demonstrated experience and expertise in applying object-oriented software design patterns and principles
    Working knowledge of C, C#, Java, C++, Python, or Ruby
    Demonstrated experience and expertise with common software development practices, including:
        Agile/Scrum or similar methodologies
        Version control and continuous integration
        Effective and comprehensive testing/TDD strategies
        Cross-platform implementations (Linux and Windows)
    Strong background in computing and the ability to demonstrate capabilities in object-oriented design and development
    Must be eligible to work in the US and able to obtain a US government security clearance
    Occasional travel may be required

Please indicate in your application materials if you have any of the following:

    Internship or similar work experience
    Experience with image processing techniques and frameworks
    Experience with scientific computing and computational geometry
    Experience with high-performance computing architectures and frameworks
    Experience integrating systems with different components using various technologies
    Experience designing modular software and communicating/refining designs independently or with a team through whiteboarding, diagrams, UML, etc.
    Experience with Docker, Kubernetes, and dynamic scaling
    Experience with cloud technologies and deployment (AWS, AWS GovCloud, Azure, etc.)
    Understanding of cybersecurity threats and mitigation techniques
    Experience with Atlassian management tools (JIRA, Confluence, Bitbucket, etc.)
    Experience working with remote customer teams/end users
    Experience with space systems (flight or ground)
    Willingness to learn new programming languages and development environments as needs arise – SQL, Bash, PowerShell, Python, Ruby, C#, C, Java, and others may be required on a situational basis
    Active DoD security clearance
    """
objective = rewrite_resume(description, resume)
cover_letter = write_cover_letter(description, resume)


document = Document("Bradley Chang Resume 2024-02-12.docx")
p = document.paragraphs
p[0]._p.clear()
p[0].add_run(objective)
style = document.styles["Normal"]
font = style.font
font.name = "Avenir Next"
font.size = Pt(8.5)
p[0].style = document.styles["Normal"]
document.save("updated-resume.docx")

